from random import choice

from openai import OpenAI
import os
import dotenv
from docx import Document

dotenv.load_dotenv()
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

def deepseek_api_call(input):
    response = client.chat.completions.create(
        # deepseek-chat
        # deepseek-reasoner
        model="deepseek-chat",
        max_tokens=8192,
        # response_format={
        #     "type": "json_object"
        # },
        messages=[
            {
                "role": "system",
                "content": "You are my helpful assistant"
            },
            {
                "role": "user",
                "content": f"""
                现在我要通过\n\n对我的文档进行分块Chunk用于RAG，请帮我在合适的地方增加\n\n以让我的文本更好的分块.
                输出以文本格式给出，内容只需要包含给我分块的文档，不需要给出其他信息
                {input}
                """

            },
        ],
        stream=False
    )
    # return response
    return response.choices[0].message.content

def read_docx(file_path_docx):
    doc=Document(file_path_docx)
    doc_text=""
    for paragraph in doc.paragraphs:
        doc_text=doc_text+paragraph.text+"\n\n"
    return doc_text
def save_docx(content, output_path):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_path)


def read_md(file_path_md):
    with open(file_path_md, 'r', encoding='utf-8') as file:
        md_text = file.read()
    return md_text
def save_md(content, output_path):
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(content)


file_path_docx=r"D:\Work\Job\企业助手\RagDoc\test\GH-ZD-HR-01休假管理制度A1.5.docx"
file_path_md=r"D:\Work\Job\企业助手\RagDoc\test\GH-ZD-HR-01休假管理制度A1.5.docx.md"

docx_text_from_docx=read_docx(file_path_docx)
docx_text_from_md=read_md(file_path_md)

docx_text_from_docx_processing=deepseek_api_call(docx_text_from_docx)
docx_text_from_md_processing=deepseek_api_call(docx_text_from_md)
print(f"docx_text_from_docx_processing{docx_text_from_docx_processing}\n")
print(f"docx_text_from_md_processing{docx_text_from_md_processing}\n")

save_md(docx_text_from_md_processing,"docx_processing.docx")
save_docx(docx_text_from_docx_processing,"docx_processing.md")

